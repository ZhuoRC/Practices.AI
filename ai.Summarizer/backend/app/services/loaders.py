"""Document loaders for various file formats"""
import io
from pathlib import Path
from typing import Optional
import PyPDF2
import docx
import requests
from bs4 import BeautifulSoup


class DocumentLoader:
    """Base class for document loaders"""

    @staticmethod
    def load_text(file_path: Path) -> str:
        """Load plain text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def load_text_from_bytes(content: bytes, encoding: str = 'utf-8') -> str:
        """Load text from bytes"""
        return content.decode(encoding)

    @staticmethod
    def load_pdf(file_path: Path) -> str:
        """Load PDF file and extract text"""
        text = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return '\n'.join(text)

    @staticmethod
    def load_pdf_from_bytes(content: bytes) -> str:
        """Load PDF from bytes and extract text"""
        text = []
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return '\n'.join(text)

    @staticmethod
    def load_docx(file_path: Path) -> str:
        """Load DOCX file and extract text"""
        doc = docx.Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        return '\n'.join(text)

    @staticmethod
    def load_docx_from_bytes(content: bytes) -> str:
        """Load DOCX from bytes and extract text"""
        doc_file = io.BytesIO(content)
        doc = docx.Document(doc_file)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        return '\n'.join(text)

    @staticmethod
    def load_webpage(url: str, timeout: int = 30) -> str:
        """Load webpage and extract main text content"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=timeout)
            response.raise_for_status()
            response.encoding = response.apparent_encoding

            soup = BeautifulSoup(response.text, 'lxml')

            # Remove script and style elements
            for element in soup(['script', 'style', 'nav', 'header', 'footer', 'aside']):
                element.decompose()

            # Get text from main content areas
            main_content = soup.find('main') or soup.find('article') or soup.find('body')
            if main_content:
                text = main_content.get_text(separator='\n', strip=True)
            else:
                text = soup.get_text(separator='\n', strip=True)

            # Clean up excessive whitespace
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            return '\n'.join(lines)

        except Exception as e:
            raise ValueError(f"Failed to load webpage: {str(e)}")

    @staticmethod
    def load_document(file_path: Path, file_type: Optional[str] = None) -> str:
        """
        Load document based on file extension

        Args:
            file_path: Path to the document
            file_type: Optional file type override (pdf, txt, docx)

        Returns:
            Extracted text content
        """
        if file_type is None:
            file_type = file_path.suffix.lower()

        if file_type in ['.txt', '.md', '.text']:
            return DocumentLoader.load_text(file_path)
        elif file_type == '.pdf':
            return DocumentLoader.load_pdf(file_path)
        elif file_type in ['.docx', '.doc']:
            return DocumentLoader.load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def load_document_from_bytes(content: bytes, file_type: str) -> str:
        """
        Load document from bytes based on file type

        Args:
            content: File content as bytes
            file_type: File extension (e.g., '.pdf', '.txt', '.docx')

        Returns:
            Extracted text content
        """
        file_type = file_type.lower()

        if file_type in ['.txt', '.md', '.text']:
            return DocumentLoader.load_text_from_bytes(content)
        elif file_type == '.pdf':
            return DocumentLoader.load_pdf_from_bytes(content)
        elif file_type in ['.docx', '.doc']:
            return DocumentLoader.load_docx_from_bytes(content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
