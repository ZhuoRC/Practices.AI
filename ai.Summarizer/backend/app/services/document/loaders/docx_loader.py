"""DOCX document loader"""
import io
from pathlib import Path
import docx


class DOCXLoader:
    """Loader for DOCX documents"""

    @staticmethod
    def load_from_path(file_path: Path) -> str:
        """Load DOCX file and extract text"""
        doc = docx.Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        return '\n'.join(text)

    @staticmethod
    def load_from_bytes(content: bytes) -> str:
        """Load DOCX from bytes and extract text"""
        doc_file = io.BytesIO(content)
        doc = docx.Document(doc_file)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        return '\n'.join(text)
